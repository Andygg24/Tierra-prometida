"""
carta_temp.py — Edita CARTA DE TEMPERATURA .docx
Recibe JSON por argv[1], escribe el .docx editado a stdout (binario).
"""
import sys, json, os, tempfile, io
from docx import Document
from datetime import datetime

data   = json.load(io.TextIOWrapper(sys.stdin.buffer, encoding="utf-8-sig"))
fecha  = data.get("fecha", "")          # "YYYY-MM-DD"
booking    = data.get("booking", "")
motonave   = data.get("motonave", "")
naviera    = data.get("naviera", "")
contenedor = data.get("contenedor", "")

# ── Formatear fecha en inglés "May 23, 2026" ─────────────────────────
MONTHS = ["January","February","March","April","May","June",
          "July","August","September","October","November","December"]
try:
    dt = datetime.strptime(fecha, "%Y-%m-%d")
    fecha_fmt = f"{MONTHS[dt.month - 1]} {dt.day}, {dt.year}"
except Exception:
    fecha_fmt = fecha

# ── Ruta del molde ────────────────────────────────────────────────────
BASE   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MOLDES = os.path.join(os.path.dirname(BASE), "Moldes")
TMPL   = os.path.join(MOLDES, "CARTA DE TEMPERATURA LIMON Booking 270711823.docx")

doc = Document(TMPL)

# ── Helper: reemplaza texto completo de un párrafo conservando formato ──
def reemplaza_parrafo(para, nuevo_texto):
    """Pone nuevo_texto en el run[0] y vacía el resto."""
    if not para.runs:
        return
    para.runs[0].text = nuevo_texto
    for r in para.runs[1:]:
        r.text = ""

# ── Helper: reemplaza substring dentro de runs de un párrafo ──────────
def reemplaza_en_runs(para, viejo, nuevo):
    for run in para.runs:
        if viejo in run.text:
            run.text = run.text.replace(viejo, nuevo)
            return True
    return False

# ── P02: "Barranquilla, May 13, 2026"
# La fecha está fragmentada en varios runs:
#   run0: "Barranquilla, "  ← conservar
#   run1: "May "            ← reemplazar con fecha completa
#   run2..N: "13", ",", " 202", "6" ← vaciar
para_fecha = doc.paragraphs[2]
runs_fecha  = para_fecha.runs
if len(runs_fecha) >= 2:
    runs_fecha[1].text = fecha_fmt
    for i in range(2, len(runs_fecha)):
        runs_fecha[i].text = ""
else:
    # Fallback: reemplazar todo el párrafo
    reemplaza_parrafo(para_fecha, f"Barranquilla, {fecha_fmt}")

# ── P07: Motonave ─────────────────────────────────────────────────────
if motonave:
    reemplaza_parrafo(doc.paragraphs[7], motonave)

# ── P08: "Booking 270711823" — reemplazar solo el número ─────────────
if booking:
    if not reemplaza_en_runs(doc.paragraphs[8], "270711823", booking):
        # Fallback: el número puede estar en cualquier run del párrafo
        for run in doc.paragraphs[8].runs:
            run_text = run.text.strip()
            if run_text.isdigit() or run_text == "":
                continue
            if "Booking" not in run_text:
                run.text = booking
                break

# ── P09: Naviera ──────────────────────────────────────────────────────
if naviera:
    reemplaza_parrafo(doc.paragraphs[9], naviera)

# ── P16: número de contenedor en la oración larga ────────────────────
if contenedor:
    if not reemplaza_en_runs(doc.paragraphs[16], "MNBU4355023", contenedor):
        # Fallback: buscar en todos los runs del documento
        for para in doc.paragraphs:
            if reemplaza_en_runs(para, "MNBU4355023", contenedor):
                break

# ── Guardar a bytes y enviar por stdout ───────────────────────────────
tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
tmp.close()
doc.save(tmp.name)

with open(tmp.name, "rb") as f:
    sys.stdout.buffer.write(f.read())

try:
    os.unlink(tmp.name)
except Exception:
    pass
