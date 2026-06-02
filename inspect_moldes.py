import xlrd
import openpyxl
from docx import Document

MOLDES = r"C:\Users\Andy Garcia\.claude\tierra-prometida\Moldes"

# ── 1. CARTA DE TEMPERATURA (.docx) ──────────────────────────────────────
print("\n" + "="*60)
print("CARTA DE TEMPERATURA (.docx)")
print("="*60)
doc = Document(rf"{MOLDES}\CARTA DE TEMPERATURA LIMON Booking 270711823.docx")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"  P{i:02d}: {para.text!r}")
for i, table in enumerate(doc.tables):
    print(f"\n  TABLE {i}:")
    for j, row in enumerate(table.rows):
        for k, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"    [{j},{k}]: {cell.text!r}")

# ── 2. PROFORMA (.xlsx) ───────────────────────────────────────────────────
print("\n" + "="*60)
print("PROFORMA (.xlsx)")
print("="*60)
wb = openpyxl.load_workbook(rf"{MOLDES}\PROFORMA PRINCESSES KINGDOM 706.xlsx")
for shname in wb.sheetnames:
    ws = wb[shname]
    print(f"\n  -- Hoja: {shname} --")
    for row in ws.iter_rows():
        for cell in row:
            if cell.value not in (None, ""):
                print(f"    {cell.coordinate}: {repr(cell.value)}")

# ── 3. ISF (.xls) ─────────────────────────────────────────────────────────
print("\n" + "="*60)
print("ISF TEMPLATE (.xls)")
print("="*60)
wb2 = xlrd.open_workbook(rf"{MOLDES}\ISF Template - JKFE AND PRINCESSES KINGDOM 270711823.xls", formatting_info=True)
for sheet in wb2.sheets():
    print(f"\n  -- Hoja: {sheet.name} ({sheet.nrows}r x {sheet.ncols}c) --")
    for r in range(sheet.nrows):
        for c in range(sheet.ncols):
            cell = sheet.cell(r, c)
            if cell.value not in ("", None, 0.0):
                col_letter = chr(ord('A') + c) if c < 26 else chr(ord('A') + c//26 - 1) + chr(ord('A') + c%26)
                print(f"    {col_letter}{r+1}: {repr(cell.value)} (ctype={cell.ctype})")

print("\nDone.")
